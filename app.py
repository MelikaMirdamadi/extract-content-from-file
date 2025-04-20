# -*- coding: utf-8 -*-
import PyPDF2
from docx import Document
import json
import re
import os
import sys
from typing import List, Dict, Tuple

# Fix console encoding for Windows
if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

def clean_text(text: str) -> str:
    """Clean text from invalid XML characters and normalize spaces"""
    if not text:
        return ""
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
    text = text.replace('\x0b', ' ').replace('\x0c', ' ')
    text = text.replace('\u0643', '\u06a9')  # ك to ک
    text = text.replace('\u0649', '\u06cc')  # ى to ی
    return text.strip()

def extract_text_by_page(pdf_path: str) -> List[Tuple[int, str]]:
    """Extract text from PDF by page number"""
    pages = []
    try:
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    cleaned = clean_text(page_text)
                    pages.append((i + 1, cleaned))  # page numbers are 1-based
    except Exception as e:
        print(f"Error reading PDF file: {str(e)}")
    return pages

def split_into_paragraphs(text: str) -> List[str]:
    """Group lines into paragraphs separated by empty lines"""
    lines = text.splitlines()
    paragraphs = []
    current = []

    for line in lines:
        stripped = line.strip()
        if stripped == '':
            if current:
                paragraphs.append(' '.join(current))
                current = []
        else:
            current.append(stripped)

    if current:
        paragraphs.append(' '.join(current))

    return paragraphs

def find_backlash_paragraphs_with_pages(pages: List[Tuple[int, str]]) -> List[Dict]:
    """Find all paragraphs with the word 'backlash', including page number"""
    results = []
    for page_num, text in pages:
        paragraphs = split_into_paragraphs(text)
        for i, para in enumerate(paragraphs, 1):
            if re.search(r'\bbacklash\b', para, re.IGNORECASE):
                sentences = re.split(r'(?<=[.!?]) +', para)
                relevant_sentences = [clean_text(s) for s in sentences
                                      if re.search(r'\bbacklash\b', s, re.IGNORECASE)]
                results.append({
                    "id": f"{page_num}.{i}",
                    "page": page_num,
                    "paragraph": clean_text(para),
                    "sentences": relevant_sentences,
                    "full_paragraph": clean_text(para)
                })
    return results

def save_to_word(occurrences: List[Dict], output_path: str):
    doc = Document()
    doc.add_heading('گزارش پاراگراف‌های حاوی کلمه Backlash', 0)

    if not occurrences:
        doc.add_paragraph("هیچ پاراگرافی حاوی کلمه 'backlash' در سند پیدا نشد.")
        doc.save(output_path)
        return

    doc.add_paragraph(f"تعداد {len(occurrences)} پاراگراف حاوی کلمه 'backlash' پیدا شد:\n")

    for item in occurrences:
        doc.add_paragraph(f"پاراگراف صفحه {item['page']}، شماره {item['id']}:", style='Heading 2')
        doc.add_paragraph(item['paragraph'])
        doc.add_paragraph("جملات مرتبط:", style='Heading 3')
        for sent in item['sentences']:
            doc.add_paragraph(f"- {sent}")
        doc.add_paragraph()

    try:
        doc.save(output_path)
    except Exception as e:
        print(f"Error saving Word file: {str(e)}")
        raise

def save_to_json(occurrences: List[Dict], output_path: str):
    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump({
                "search_term": "backlash",
                "total_paragraphs": len(occurrences),
                "paragraphs": [ {
                    "id": item["id"],
                    "page": item["page"],
                    "full_paragraph": item["full_paragraph"],
                    "sentences": item["sentences"]
                } for item in occurrences ]
            }, f, indent=2, ensure_ascii=False)
    except Exception as e:
        print(f"Error saving JSON file: {str(e)}")
        raise

def process_directory(input_dir: str, output_dir: str):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    for filename in os.listdir(input_dir):
        if filename.lower().endswith('.pdf'):
            pdf_path = os.path.join(input_dir, filename)
            base_name = os.path.splitext(filename)[0]
            word_output = os.path.join(output_dir, f"{base_name}_backlash_paragraphs.docx")
            json_output = os.path.join(output_dir, f"{base_name}_backlash_paragraphs.json")

            try:
                print(f"\n📄 Processing file: {filename}")
                pages = extract_text_by_page(pdf_path)
                occurrences = find_backlash_paragraphs_with_pages(pages)
                save_to_word(occurrences, word_output)
                save_to_json(occurrences, json_output)

                print(f"✅ Done! Found {len(occurrences)} matches.")
                print(f"📄 Word saved to: {word_output}")
                print(f"🗂️ JSON saved to: {json_output}")

            except Exception as e:
                print(f"❌ Failed to process {filename}: {str(e)}")

def main():
    input_dir = "./inputs"   # Replace with your input directory path
    output_dir = "./outputs" # Replace with your output directory path
    process_directory(input_dir, output_dir)

if __name__ == "__main__":
    main()
