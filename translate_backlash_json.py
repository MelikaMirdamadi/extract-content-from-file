import json
import os
from docx import Document
import sys
import subprocess

# Ensure UTF-8 for Windows terminal
if sys.stdout.encoding.lower() != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

def translate_to_persian(text):
    """Translate English text to Persian using Mistral model via Ollama."""
    try:
        # Create prompt with special instructions for technical terms
        prompt = f"""Please translate the following English text to Persian.
        Keep technical terms like 'backlash', 'gear' etc. in English:
        
        {text}
        
        Translation:"""

        # Run Ollama command
        result = subprocess.run(
            ['ollama', 'run', 'mistral', prompt],
            capture_output=True,
            text=True,
            check=True
        )

        # Get the translated text
        translated = result.stdout.strip()
        
        # Basic validation
        if not translated:
            return text
            
        return translated

    except Exception as e:
        print(f" Translation error: {str(e)[:100]}...")
        return text

def translate_backlash_json(json_input_path, output_dir):
    with open(json_input_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    output_json = data.copy()
    output_json["paragraphs"] = []

    for item in data["paragraphs"]:
        translated_paragraph = translate_to_persian(item["full_paragraph"])
        translated_sentences = [translate_to_persian(s) for s in item["sentences"]]

        item["paragraph_fa"] = translated_paragraph
        item["sentences_fa"] = translated_sentences

        output_json["paragraphs"].append(item)

    # Save JSON
    json_filename = os.path.basename(json_input_path)
    base_name = os.path.splitext(json_filename)[0]
    json_out_path = os.path.join(output_dir, f"{base_name}_translated.json")

    with open(json_out_path, 'w', encoding='utf-8') as f:
        json.dump(output_json, f, ensure_ascii=False, indent=2)

    # Save Word
    doc = Document()
    doc.add_heading("ترجمه پاراگراف‌های حاوی 'backlash'", 0)

    for item in output_json["paragraphs"]:
        doc.add_paragraph(f"پاراگراف #{item['id']}", style='Heading 2')
        doc.add_paragraph(f"English: {item['full_paragraph']}")
        doc.add_paragraph(f"Persian: {item['paragraph_fa']}")
        doc.add_paragraph("جملات:", style='Heading 3')
        for en, fa in zip(item["sentences"], item["sentences_fa"]):
            doc.add_paragraph(f"EN: {en}")
            doc.add_paragraph(f"FA: {fa}")
        doc.add_paragraph()

    word_out_path = os.path.join(output_dir, f"{base_name}_translated.docx")
    doc.save(word_out_path)

    print(f" Translations saved:\n- {json_out_path}\n- {word_out_path}")

# Example usage
if __name__ == "__main__":
    input_json = "./outputs/Cylindrical_Gears_Calculation_Materials_Manufacturing_2016_Linke_backlash_paragraphs.json"  # <--- Set your path
    output_folder = "./translated_outputs"
    os.makedirs(output_folder, exist_ok=True)
    translate_backlash_json(input_json, output_folder)
